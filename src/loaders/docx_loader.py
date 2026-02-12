from pathlib import Path
from typing import Optional
from docx import Document as DocxDocument

from src.loaders.base_loader import BaseLoader
from src.utils.document import Document

"""
DOCXLoader - Load Microsoft Word documents
Extracts text from .docx files preserving paragraph structure
"""
"""
Load a Microsoft Word (.docx) file and extract text content
Preserves paragraph structure and can optionally extract
text from tables
Attributes:
    file_path: Path to the Word document
    include_tables: Whether to extract table content
Example:
    >>> loader = DOCXLoader("policy.docx")
    >>> docs = loader.load()
    >>> print(docs[0].page_content[:100])
"""
class DOCXLoader(BaseLoader):
    
    """
    Initialize the DOCXLoader    
    Args:
        file_path: Path to the Word document
        include_tables: Whether to include table content
        metadata: Optional additional metadata
    """
    def __init__(
        self,
        file_path: str,
        include_tables: bool = True,
        metadata: Optional[dict] = None
    ):
        self.file_path = Path(file_path)
        self.include_tables = include_tables
        self.extra_metadata = metadata or {}
    
    """
    Load the Word document and extract text
    Returns:
        A list containing a single Document with all text content
    Raises:
        FileNotFoundError: If the file doesn't exist
    """
    def load(self) -> list[Document]:
        if not self.file_path.exists():
            raise FileNotFoundError(f"File not found: {self.file_path}")
        
        doc = DocxDocument(self.file_path)
        
        """
        Extract paragraphs
        """
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        """
        Extract tables if requested
        """
        if self.include_tables:
            for table in doc.tables:
                table_text = self._extract_table(table)
                if table_text:
                    paragraphs.append(table_text)
        
        content = "\n\n".join(paragraphs)
        
        metadata = {
            "source": str(self.file_path),
            "file_name": self.file_path.name,
            "file_type": "docx",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            **self.extra_metadata
        }
        
        return [Document(page_content=content, metadata=metadata)]
    
    """
    Extract text from a table in markdown-like format
    """
    def _extract_table(self, table) -> str:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)
